from pathlib import Path

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "sample-documents" / "poc-multi-file-test"
SAME = OUT / "same-account-orion-retail"
MULTI = OUT / "multiple-accounts-portfolio"


def ensure_dirs() -> None:
    SAME.mkdir(parents=True, exist_ok=True)
    MULTI.mkdir(parents=True, exist_ok=True)


def write_text(path: Path, text: str) -> None:
    path.write_text(text.strip() + "\n", encoding="utf-8")


def make_pdf(path: Path, title: str, lines: list[str]) -> None:
    pdf = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    y = height - 72
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(72, y, title)
    y -= 28
    pdf.setFont("Helvetica", 10)
    for line in lines:
        if not line:
            y -= 10
            continue
        for segment in wrap_line(line, 92):
            if y < 72:
                pdf.showPage()
                y = height - 72
                pdf.setFont("Helvetica", 10)
            pdf.drawString(72, y, segment)
            y -= 14
    pdf.save()


def wrap_line(line: str, width: int) -> list[str]:
    words = line.split()
    lines: list[str] = []
    current: list[str] = []
    for word in words:
        if len(" ".join(current + [word])) > width and current:
            lines.append(" ".join(current))
            current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines


def make_docx(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=2)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
    doc.save(path)


def make_doc_html(path: Path, title: str, rows: list[tuple[str, str]]) -> None:
    body_rows = "\n".join(
        f"<tr><th>{label}</th><td>{value}</td></tr>" for label, value in rows
    )
    html = f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{title}</title>
  <style>
    body {{ font-family: Arial, sans-serif; }}
    h1 {{ font-size: 20px; }}
    table {{ border-collapse: collapse; width: 100%; }}
    th, td {{ border: 1px solid #999; padding: 6px; text-align: left; }}
  </style>
</head>
<body>
  <h1>{title}</h1>
  <table>{body_rows}</table>
</body>
</html>"""
    path.write_text(html, encoding="utf-8")


def make_xls_html(path: Path, title: str, headers: list[str], rows: list[list[str]]) -> None:
    header_html = "".join(f"<th>{header}</th>" for header in headers)
    row_html = "\n".join(
        "<tr>" + "".join(f"<td>{cell}</td>" for cell in row) + "</tr>" for row in rows
    )
    html = f"""<html>
<head><meta charset="utf-8"><title>{title}</title></head>
<body>
<table>
<tr><th colspan="{len(headers)}">{title}</th></tr>
<tr>{header_html}</tr>
{row_html}
</table>
</body>
</html>"""
    path.write_text(html, encoding="utf-8")


def make_xlsx(path: Path, title: str, sheets: dict[str, tuple[list[str], list[list[object]]]]) -> None:
    wb = Workbook()
    default = wb.active
    wb.remove(default)
    for sheet_name, (headers, rows) in sheets.items():
        ws = wb.create_sheet(sheet_name)
        ws.append([title])
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(1, len(headers)))
        ws["A1"].font = Font(bold=True, size=14)
        ws.append(headers)
        for cell in ws[2]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="D9EAF7")
        for row in rows:
            ws.append(row)
        for index, col in enumerate(ws.columns, start=1):
            width = max(len(str(cell.value or "")) for cell in col) + 2
            ws.column_dimensions[get_column_letter(index)].width = min(max(width, 12), 36)
    wb.save(path)


def build_same_account() -> None:
    make_pdf(
        SAME / "orion-contract-summary.pdf",
        "Orion Retail Group - Contract Summary",
        [
            "Account: Orion Retail Group",
            "Industry: Retail technology and digital commerce",
            "ARR: $1.8M",
            "Contract start: 2025-01-15",
            "Contract end / renewal: 2026-09-30",
            "Notice period: 60 days before renewal.",
            "Price terms: annual price review language exists, but no automatic uplift is guaranteed.",
            "Renewal structure: manual renewal path; no auto-renewal clause.",
            "Contract health note: commercial structure is stable but renewal confidence depends on defect closure before September 2026.",
        ],
    )

    make_docx(
        SAME / "orion-stakeholder-kyc.docx",
        "Orion Retail Group - Stakeholder and KYC Notes",
        [
            ("Executive Summary", [
                "Orion Retail Group is modernizing digital commerce, mobile checkout, and loyalty analytics.",
                "Tkxel is considered a strategic delivery partner, but procurement and CIO engagement remain incomplete.",
            ]),
            ("Stakeholders", [
                "Executive sponsor: Dana Walsh, Chief Digital Officer.",
                "Primary contact: Marcus Reed, VP Ecommerce Platforms.",
                "Procurement contact is not mapped. CIO has missed the last two steering reviews.",
            ]),
            ("Strategic Goals", [
                "Improve mobile checkout conversion, reduce payment defects, and launch loyalty analytics by Q4.",
                "Potential expansion areas include personalization, managed QA automation, and analytics support.",
            ]),
        ],
    )

    make_doc_html(
        SAME / "orion-resource-plan.doc",
        "Orion Retail Group - Resource Plan",
        [
            ("Delivery Lead", "New delivery lead joined in May and improved communication cadence."),
            ("Critical Dependency", "Payment gateway specialist is a single point of dependency."),
            ("Backup Readiness", "Backup engineer identified but has not completed checkout architecture onboarding."),
            ("Team Stability", "Core squad is stable; two sprint commitments slipped because of gateway defects."),
            ("Resource Health", "Moderate risk until specialist backup is fully ready."),
        ],
    )

    write_text(
        SAME / "orion-risk-and-delivery-log.txt",
        """
Account: Orion Retail Group
Delivery update: mobile checkout modernization is mostly on track.
Open risks: unresolved payment gateway defects remain before renewal.
Escalation: no active executive escalation, but client warned that unresolved defects must close before renewal.
Competitive note: a competing digital commerce vendor is running a discovery workshop.
Project health: two sprint commitments slipped in May; backlog is now replanned.
Risk score context: risk is moderate because churn risk depends on defect closure and competitor activity.
""",
    )

    write_text(
        SAME / "orion-meeting-notes.md",
        """
# Orion Retail Group Meeting Notes

## CSAT
Latest CSAT is 4.1 out of 5. Customer sentiment is generally satisfied with roadmap alignment.

## Relationship
Monthly steering calls happen with the VP Ecommerce Platforms. Quarterly executive reviews happen with the CDO.

## Whitespace
Client is exploring loyalty analytics, personalization, and managed QA automation.

## Missing Information
Detailed procurement ownership and CIO sponsorship are not yet confirmed.
""",
    )

    make_xls_html(
        SAME / "orion-financial-ledger.xls",
        "Orion Retail Group - Financial Ledger",
        ["Metric", "Value", "Evidence"],
        [
            ["ARR", "$1.8M", "Contract summary"],
            ["Invoice status", "Current", "No overdue invoices as of latest finance review"],
            ["Billing risk", "Low", "Billing matches current SOW"],
            ["Expansion potential", "Moderate", "Analytics and QA automation under discussion"],
        ],
    )

    make_xlsx(
        SAME / "orion-csat-project-dashboard.xlsx",
        "Orion Retail Group - CSAT and Project Dashboard",
        {
            "CSAT": (
                ["Period", "CSAT", "NPS", "Comment"],
                [
                    ["2026-Q1", 4.0, 38, "Stable delivery sentiment"],
                    ["2026-Q2", 4.1, 41, "Improved communication after new delivery lead"],
                ],
            ),
            "Project": (
                ["Workstream", "Status", "Risk", "Owner"],
                [
                    ["Mobile checkout", "Watchlist", "Gateway defects", "Tkxel delivery lead"],
                    ["Loyalty analytics", "Opportunity", "Discovery pending", "KAM"],
                    ["Managed QA automation", "Opportunity", "Business case pending", "KAM"],
                ],
            ),
        },
    )


def build_multiple_accounts() -> None:
    make_pdf(
        MULTI / "portfolio-account-briefs.pdf",
        "Portfolio Account Briefs - Three Accounts",
        [
            "Account: Apex Manufacturing Cloud | ARR: $950K | Region: EMEA | Contract end: 2026-05-31",
            "Apex has low executive engagement, unresolved ERP integration defects, and a competitor-led assessment in progress.",
            "",
            "Account: BrightBank Digital | ARR: $2.4M | Region: North America | Contract end: 2027-02-15",
            "BrightBank has strong sponsor engagement, current invoices, and active expansion interest in fraud analytics.",
            "",
            "Account: Helio Health Systems | ARR: $1.2M | Region: APAC | Contract end: 2026-08-20",
            "Helio has delayed payments, limited backup resources, and rising support tickets around patient portal releases.",
        ],
    )

    make_docx(
        MULTI / "portfolio-renewal-notes.docx",
        "Portfolio Renewal Notes",
        [
            ("Apex Manufacturing Cloud", [
                "Renewal is at risk because the CFO sponsor has not attended recent reviews.",
                "Contract has a 30-day notice period and no price uplift language.",
            ]),
            ("BrightBank Digital", [
                "Renewal confidence is high. Executive sponsor attends quarterly reviews.",
                "Potential cross-sell: fraud analytics, cloud cost optimization, and managed platform support.",
            ]),
            ("Helio Health Systems", [
                "Renewal requires corrective action on tickets, payment aging, and resource continuity.",
                "Client requested a recovery plan before the next steering committee.",
            ]),
        ],
    )

    make_doc_html(
        MULTI / "portfolio-resource-notes.doc",
        "Portfolio Resource Notes",
        [
            ("Apex Manufacturing Cloud", "Two key engineers are rotating off; backup knowledge transfer is incomplete."),
            ("BrightBank Digital", "Team is stable with documented backups for fraud analytics and platform support."),
            ("Helio Health Systems", "Patient portal lead is overloaded; no fully onboarded backup exists."),
        ],
    )

    write_text(
        MULTI / "portfolio-signals.txt",
        """
CHURN_RISK: Apex Manufacturing Cloud has a competitor assessment in progress.
CONTRACT_EXPIRY: Apex renewal date is 2026-05-31 with only 30-day notice.
UPSELL_OPPORTUNITY: BrightBank Digital asked for fraud analytics discovery.
TICKET_SPIKE: Helio Health Systems opened 17 patient portal tickets in the last 30 days.
NPS_DECLINE: Helio NPS moved from 21 to 8 quarter over quarter.
""",
    )

    write_text(
        MULTI / "portfolio-kyc-notes.md",
        """
# Portfolio KYC Notes

## Apex Manufacturing Cloud
Business model: industrial manufacturing and connected ERP operations.
Risk factors: low executive engagement, unresolved integration defects, competitor assessment.

## BrightBank Digital
Business model: digital banking, fraud monitoring, and customer onboarding workflows.
Strategic goals: improve fraud analytics, stabilize platform costs, expand managed support.

## Helio Health Systems
Business model: hospital network with patient portal and care coordination tools.
Risk factors: delayed payments, ticket spike, resource dependency, declining NPS.
""",
    )

    make_xls_html(
        MULTI / "portfolio-finance-and-csat.xls",
        "Portfolio Finance and CSAT",
        ["Account", "ARR", "Invoice Status", "CSAT", "Financial Risk"],
        [
            ["Apex Manufacturing Cloud", "$950K", "Current", "3.0/5", "Medium"],
            ["BrightBank Digital", "$2.4M", "Current", "4.6/5", "Low"],
            ["Helio Health Systems", "$1.2M", "45 days overdue", "2.7/5", "High"],
        ],
    )

    make_xlsx(
        MULTI / "portfolio-account-metrics.xlsx",
        "Portfolio Account Metrics",
        {
            "Scores": (
                ["Account", "Relationship", "Contract", "CSAT", "Risk", "Resource", "Project", "Financial", "Whitespace"],
                [
                    ["Apex Manufacturing Cloud", 2.1, 2.0, 3.0, 2.0, 2.4, 2.6, 3.0, 2.2],
                    ["BrightBank Digital", 4.5, 4.2, 4.6, 4.4, 4.1, 4.3, 4.7, 4.5],
                    ["Helio Health Systems", 2.8, 2.9, 2.7, 2.2, 1.9, 2.5, 1.8, 2.4],
                ],
            ),
            "Signals": (
                ["Account", "Signal", "Severity", "Evidence"],
                [
                    ["Apex Manufacturing Cloud", "CHURN_RISK", "CRITICAL", "Competitor assessment active"],
                    ["BrightBank Digital", "UPSELL_OPPORTUNITY", "INFO", "Fraud analytics discovery requested"],
                    ["Helio Health Systems", "TICKET_SPIKE", "WARNING", "17 tickets in 30 days"],
                ],
            ),
        },
    )


def build_readme() -> None:
    write_text(
        OUT / "README.md",
        """
# POC Multi-File Test Pack

Use `same-account-orion-retail` to test multiple files that should be combined into one account extraction.

Use `multiple-accounts-portfolio` to test files that mention several accounts. The POC will still produce one consolidated extraction, so this folder is useful for seeing whether the AI asks for clarification or focuses on the dominant account.

Included formats:
- PDF
- DOC (Word-compatible HTML .doc sample)
- DOCX
- TXT
- MD
- XLS (Excel-compatible HTML .xls sample)
- XLSX
""",
    )


def main() -> None:
    ensure_dirs()
    build_same_account()
    build_multiple_accounts()
    build_readme()
    print(f"Generated sample files in: {OUT}")


if __name__ == "__main__":
    main()
